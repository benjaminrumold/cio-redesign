#!/usr/bin/env python3
"""Create CiO Redesign InfoSheet as a professional Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Colors
NAVY = RGBColor(0x0F, 0x1B, 0x3D)
ROSE = RGBColor(0xB4, 0x43, 0x6C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_NAVY_HEX = "1a2d5a"  # lighter navy for table headers
LIGHT_GRAY_HEX = "f5f5f5"

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Default Style ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Heading Styles ──
for level, size, space_before, space_after in [
    ('Heading 1', 14, 24, 12),
    ('Heading 2', 12, 18, 8),
    ('Heading 3', 11, 12, 6),
]:
    h = doc.styles[level]
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = NAVY
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)

# ── List Bullet Style ──
list_style = doc.styles['List Bullet']
list_style.font.name = 'Calibri'
list_style.font.size = Pt(11)
list_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ════════════════════════════════════════════════════════════

def add_accent_line(doc):
    """Add a thin rose-colored accent line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Add bottom border to paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="b4436c"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="cccccc", size="4"):
    """Set thin borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def format_header_cell(cell, text):
    """Format a table header cell with navy background and white text."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_cell_shading(cell, LIGHT_NAVY_HEX)
    set_cell_borders(cell, "0f1b3d", "4")


def format_data_cell(cell, text, row_idx=0):
    """Format a table data cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    if row_idx % 2 == 1:
        set_cell_shading(cell, LIGHT_GRAY_HEX)
    set_cell_borders(cell, "cccccc", "4")


def create_table(doc, headers, rows, col_widths=None):
    """Create a professionally formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    for i, header in enumerate(headers):
        format_header_cell(table.rows[0].cells[i], header)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            format_data_cell(table.rows[r_idx + 1].cells[c_idx], cell_text, r_idx)

    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    return table


def add_section_heading(doc, number, title):
    """Add a numbered section heading."""
    h = doc.add_heading(level=1)
    run = h.add_run(f"{number}. {title}")
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    add_accent_line(doc)


def add_sub_heading(doc, title):
    """Add a sub-heading."""
    h = doc.add_heading(title, level=2)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = 'Calibri'


def add_body(doc, text):
    """Add body text paragraph."""
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p


def add_bullet(doc, text):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    return p


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ════════════════════════════════════════════════════════════
# PAGE NUMBERS IN FOOTER
# ════════════════════════════════════════════════════════════

section = doc.sections[0]
footer = section.footer
footer.is_linked_to_previous = False
p_footer = footer.paragraphs[0]
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_footer.add_run()
run.font.name = 'Calibri'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
# Add page number field
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run._r.append(instrText)
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run._r.append(fldChar2)


# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════

# Spacer
for _ in range(6):
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Christen in der Osteopathie")
run.font.name = 'Calibri'
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = NAVY

# Accent line under title
add_accent_line(doc)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
run = p.add_run("Website-Redesign 2026: Analyse, Zielgruppe & Strategie")
run.font.name = 'Calibri'
run.font.size = Pt(16)
run.font.color.rgb = ROSE

# Spacer
doc.add_paragraph()

# Date
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Mai 2026")
run.font.name = 'Calibri'
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Author
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Erstellt von: Benjamin Rumold")
run.font.name = 'Calibri'
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Spacer
for _ in range(4):
    doc.add_paragraph()

# Confidentiality notice
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run("Vertraulich — Nur für internen Gebrauch")
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

add_page_break(doc)


# ════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════

add_section_heading(doc, "1", "EXECUTIVE SUMMARY")

add_body(doc, (
    "Die Website christen-in-der-osteopathie.de ist die zentrale Online-Präsenz des Netzwerks "
    "„Christen in der Osteopathie“ (CiO). Die aktuelle Seite basiert auf einem veralteten "
    "WordPress-Theme von 2017 und weist erhebliche Defizite in den Bereichen Mobile-Optimierung, "
    "Suchmaschinenoptimierung (SEO), Barrierefreiheit und Nutzerführung auf."
))

add_body(doc, (
    "Dieses Dokument analysiert den Ist-Zustand, definiert die Zielgruppe anhand fundierter Daten, "
    "identifiziert rechtliche Anforderungen und leitet daraus konkrete Redesign-Maßnahmen ab. "
    "Ziel ist eine moderne, professionelle und DSGVO-konforme Website, die das Netzwerk zeitgemäß "
    "repräsentiert und neue Mitglieder sowie Patienten effektiv anspricht."
))

add_page_break(doc)


# ════════════════════════════════════════════════════════════
# 2. IST-ANALYSE
# ════════════════════════════════════════════════════════════

add_section_heading(doc, "2", "IST-ANALYSE: AKTUELLE WEBSITE")

ist_rows = [
    ("Viewport Meta-Tag", "❌ Fehlt komplett — keine mobile Optimierung"),
    ("Meta-Description", "❌ Keine vorhanden — unsichtbar für Google"),
    ("Responsive Design", "❌ Nicht implementiert — Desktop-only Layout"),
    ("SSL-Zertifikat", "✅ HTTPS aktiv"),
    ("Ladezeit", "⚠️ Nicht optimiert (kein Lazy Loading, keine Bildkompression)"),
    ("Therapeutensuche", "❌ Nur statische Liste ohne Filter- oder Suchfunktion"),
    ("Veranstaltungen", "⚠️ Preise erst auf Detailseite sichtbar"),
    ("Barrierefreiheit", "❌ Keine ARIA-Labels, kein Fokus-Management"),
    ("WordPress-Theme", "❌ Veraltet (2017), keine Updates"),
    ("Content-Struktur", "⚠️ Unstrukturiert, keine klare Nutzerführung"),
]

create_table(doc, ["Kriterium", "Bewertung"], ist_rows, col_widths=[2.2, 4.3])

doc.add_paragraph()
add_body(doc, (
    "Die Website erfüllt damit weder aktuelle technische Standards noch die ab Juni 2025 "
    "geltenden Anforderungen des Barrierefreiheitsstärkungsgesetzes (BFSG)."
))

add_page_break(doc)


# ════════════════════════════════════════════════════════════
# 3. ZIELGRUPPENANALYSE
# ════════════════════════════════════════════════════════════

add_section_heading(doc, "3", "ZIELGRUPPENANALYSE")

# 3.1
add_sub_heading(doc, "3.1 Primäre Zielgruppe: Osteopathen mit christlichem Hintergrund")

zg1_rows = [
    ("Alter", "35–55 Jahre (Schwerpunkt 40–50)"),
    ("Geschlecht", "Überwiegend weiblich (ca. 70%)"),
    ("Berufsgruppe", "Selbstständige Osteopathen, HP-Physio, Heilpraktiker"),
    ("Ausbildung", "Mind. 4 Jahre (1.350+ Stunden), akademisch orientiert"),
    ("Digitalkompetenz", "Mittel bis hoch (89–97% Internetnutzung in Altersgruppe)"),
    ("Geräte-Nutzung", "Desktop primär (Praxis-PC), Smartphone sekundär"),
    ("Social Media", "Facebook aktiv, Instagram zunehmend, LinkedIn beruflich"),
    ("Werte", "Glaube, Fachkompetenz, Gemeinschaft, evidenzbasierte Medizin"),
    ("Motivation", "Fachliche Fortbildung + geistlicher Austausch"),
]

create_table(doc, ["Merkmal", "Detail"], zg1_rows, col_widths=[2.0, 4.5])
doc.add_paragraph()

# 3.2
add_sub_heading(doc, "3.2 Sekundäre Zielgruppe: Patienten")

zg2_rows = [
    ("Alter", "30–65 Jahre"),
    ("Suchverhalten", "„Osteopath + [Stadt]“, „christlicher Osteopath“"),
    ("Bedürfnis", "Therapeut finden, der Werte teilt"),
    ("Geräte", "Überwiegend Smartphone (Google-Suche)"),
    ("Entscheidungsfaktor", "Nähe, Qualifikation, Vertrauenssignale"),
]

create_table(doc, ["Merkmal", "Detail"], zg2_rows, col_widths=[2.0, 4.5])
doc.add_paragraph()

# 3.3
add_sub_heading(doc, "3.3 Marktkontext")

add_body(doc, (
    "In Deutschland praktizieren ca. 7.000–10.000 Osteopathen. Der VOD (Verband der Osteopathen "
    "Deutschland) hat über 6.500 Mitglieder, der BVO ca. 1.400. CiO bedient eine spezifische Nische: "
    "die Verbindung von osteopathischer Fachkompetenz mit christlichem Glauben."
))

add_body(doc, "Vergleichbare christliche Gesundheitsnetzwerke:")
add_bullet(doc, "Christen im Gesundheitswesen (CiG) — größtes Netzwerk, breiter Ansatz")
add_bullet(doc, "C-STAB — Christliche Heilpraktiker")
add_bullet(doc, "DGCN — Christliche Naturheilkunde")
add_bullet(doc, "ACM — Akademie für christliche Medizin")

add_page_break(doc)


# ════════════════════════════════════════════════════════════
# 4. BENCHMARKS & WETTBEWERB
# ════════════════════════════════════════════════════════════

add_section_heading(doc, "4", "BENCHMARKS & WETTBEWERB")

bench_rows = [
    ("osteopathie.de (VOD)", "Therapeutensuche mit PLZ, professionelles Design", "Kein religiöser Bezug, sehr funktional"),
    ("bv-osteopathie.de (BVO)", "Moderne Optik, gute Struktur", "Wenig Community-Fokus"),
    ("christenimgesundheitswesen.de", "Starke Community, Events", "Veraltetes Design, langsam"),
    ("dgcn.de", "Nischen-Positionierung", "Minimal-Design, wenig Funktionalität"),
]

create_table(doc, ["Website", "Stärken", "Schwächen"], bench_rows, col_widths=[2.3, 2.2, 2.0])

add_page_break(doc)


# ════════════════════════════════════════════════════════════
# 5. RECHTLICHE ANFORDERUNGEN
# ════════════════════════════════════════════════════════════

add_section_heading(doc, "5", "RECHTLICHE ANFORDERUNGEN")

# 5.1
add_sub_heading(doc, "5.1 DSGVO (Datenschutz-Grundverordnung)")
add_bullet(doc, "Datenschutzerklärung nach Art. 13 DSGVO")
add_bullet(doc, "Cookie-Consent-Banner mit Opt-in")
add_bullet(doc, "Auftragsverarbeitungsverträge (AVV) für alle Drittanbieter")
add_bullet(doc, "SSL-Verschlüsselung (bereits vorhanden)")
add_bullet(doc, "Recht auf Auskunft, Löschung, Datenportabilität")

# 5.2
add_sub_heading(doc, "5.2 Heilmittelwerbegesetz (HWG)")
add_bullet(doc, "Keine Heilversprechen oder Erfolgsgarantien")
add_bullet(doc, "Keine irreführenden Vorher-Nachher-Darstellungen")
add_bullet(doc, "Sachliche Darstellung der Osteopathie")
add_bullet(doc, "Therapeutenliste ohne werbliche Übertreibung")

# 5.3
add_sub_heading(doc, "5.3 Barrierefreiheitsstärkungsgesetz (BFSG)")
add_body(doc, "Ab 28. Juni 2025 gelten erweiterte Barrierefreiheitsanforderungen:")
add_bullet(doc, "WCAG 2.1 Level AA als Mindeststandard")
add_bullet(doc, "Tastaturbedienbarkeit aller Funktionen")
add_bullet(doc, "Screenreader-Kompatibilität (ARIA-Labels)")
add_bullet(doc, "Ausreichende Kontrastverhältnisse (min. 4.5:1)")
add_bullet(doc, "Skalierbare Schriften (bis 200% ohne Funktionsverlust)")
add_bullet(doc, "Alternative Texte für alle Bilder")

# 5.4
add_sub_heading(doc, "5.4 Impressumspflicht (TMG/DDG)")
add_bullet(doc, "Vollständiges Impressum nach § 5 DDG")
add_bullet(doc, "Angabe der zuständigen Aufsichtsbehörde")
add_bullet(doc, "Berufsbezeichnung und Kammerzugehörigkeit")

add_page_break(doc)


# ════════════════════════════════════════════════════════════
# 6. UX-EMPFEHLUNGEN
# ════════════════════════════════════════════════════════════

add_section_heading(doc, "6", "UX-EMPFEHLUNGEN FÜR DIE ZIELGRUPPE")

ux_rows = [
    ("Mobile-First Design", "89-97% der Zielgruppe nutzt das Internet, zunehmend mobil", "Kritisch"),
    ("Große, lesbare Schriften (min. 16px)", "Altersgruppe 40-55, Lesbarkeit am Bildschirm", "Hoch"),
    ("Klare Navigation (max. 6 Hauptpunkte)", "Mittlere Digitalkompetenz, schnelle Orientierung", "Hoch"),
    ("Therapeutensuche mit PLZ-Filter", "Kernfunktion für Patienten, Benchmark bei VOD", "Kritisch"),
    ("Veranstaltungskalender mit Preisen", "Transparenz als Vertrauensfaktor", "Hoch"),
    ("Kontaktformular statt nur E-Mail", "Niedrigschwelliger Erstkontakt", "Mittel"),
    ("Testimonials / Erfahrungsberichte", "Social Proof für beide Zielgruppen", "Mittel"),
    ("Blog / Aktuelles", "SEO-Boost, Wiederkehrbesucher", "Niedrig"),
]

create_table(doc, ["Empfehlung", "Begründung", "Priorität"], ux_rows, col_widths=[2.3, 3.0, 1.2])

add_page_break(doc)


# ════════════════════════════════════════════════════════════
# 7. STÄRKEN DER AKTUELLEN WEBSITE
# ════════════════════════════════════════════════════════════

add_section_heading(doc, "7", "STÄRKEN DER AKTUELLEN WEBSITE")

staerken = [
    "Einzigartige Positionierung: Einziges Netzwerk für christliche Osteopathen im DACH-Raum",
    "Authentischer Content: Persönliche Ansprache, klare Wertebasis",
    "Starkes Fortbildungsprogramm: 6 Seminare + 2 Netzwerktreffen in 2026/2027",
    "Zertifizierung: VOD- und BVO-anerkannte Fortbildungspunkte",
    "Etabliertes Netzwerk: 40+ Therapeuten aus Deutschland, Österreich und der Schweiz",
    "SSL-Verschlüsselung: Grundlegende Sicherheit vorhanden",
    "Klare Abgrenzung: Explizite Distanzierung von Esoterik, naturwissenschaftliche Basis",
]
for s in staerken:
    add_bullet(doc, s)

add_page_break(doc)


# ════════════════════════════════════════════════════════════
# 8. REDESIGN-MASSNAHMEN
# ════════════════════════════════════════════════════════════

add_section_heading(doc, "8", "REDESIGN-MASSNAHMEN")

# 8.1
add_sub_heading(doc, "8.1 Technische Maßnahmen")
for item in [
    "Viewport Meta-Tag implementieren",
    "Responsive Layout (Mobile-First)",
    "Bildoptimierung (WebP, Lazy Loading)",
    "Strukturierte Daten (Schema.org: MedicalOrganization, Event)",
    "Performance-Optimierung (Core Web Vitals)",
    "ARIA-Labels und Fokus-Management",
    "Cookie-Consent-Integration",
]:
    add_bullet(doc, item)

# 8.2
add_sub_heading(doc, "8.2 Inhaltliche Maßnahmen")
for item in [
    "SEO-optimierte Meta-Descriptions für alle Seiten",
    "Therapeutensuche mit PLZ-Filter und Kartenansicht",
    "Veranstaltungsübersicht mit allen Preisen und Terminen",
    "Klare Call-to-Actions (Anmeldung, Kontakt, Mitglied werden)",
    "FAQ-Bereich für häufige Fragen",
    "Download-Bereich (Flyer, Anmeldeformulare)",
]:
    add_bullet(doc, item)

# 8.3
add_sub_heading(doc, "8.3 Design-Maßnahmen")
for item in [
    "Modernes Farbschema: Navy (#0f1b3d) + Rose (#b4436c) + Gold (#c9a96e)",
    "Professionelle Typografie: Cormorant Garamond + DM Sans",
    "Scroll-Animationen (Intersection Observer)",
    "Hero-Section mit animiertem Hintergrund",
    "Sticky Navigation mit Blur-Effekt",
    "Professionelle Bildsprache (Team, Seminare, Praxisräume)",
]:
    add_bullet(doc, item)

add_page_break(doc)


# ════════════════════════════════════════════════════════════
# 9. ZUSAMMENFASSUNG & NÄCHSTE SCHRITTE
# ════════════════════════════════════════════════════════════

add_section_heading(doc, "9", "ZUSAMMENFASSUNG & NÄCHSTE SCHRITTE")

add_body(doc, (
    "Das Redesign der CiO-Website ist nicht nur ein visuelles Update, sondern eine strategische "
    "Neuausrichtung der digitalen Präsenz. Die Kombination aus technischer Modernisierung, "
    "zielgruppengerechter Gestaltung und rechtlicher Compliance schafft die Grundlage für "
    "nachhaltiges Wachstum des Netzwerks."
))

add_sub_heading(doc, "Empfohlener Zeitplan")

timeline_rows = [
    ("Phase 1: Konzeption", "Juni 2026", "Wireframes, Content-Strategie, Designsystem"),
    ("Phase 2: Design", "Juli 2026", "UI-Design, Prototyp, Abstimmung"),
    ("Phase 3: Entwicklung", "Aug–Sep 2026", "Frontend, CMS-Integration, Therapeuten-DB"),
    ("Phase 4: Testing", "Oktober 2026", "Cross-Browser, Mobile, Barrierefreiheit, SEO"),
    ("Phase 5: Launch", "November 2026", "Go-Live, Monitoring, Optimierung"),
]

create_table(doc, ["Phase", "Zeitraum", "Maßnahmen"], timeline_rows, col_widths=[2.0, 1.5, 3.0])

doc.add_paragraph()
add_body(doc, (
    "Mit diesem Redesign positioniert sich CiO als modernes, professionelles Netzwerk, das sowohl "
    "fachliche Exzellenz als auch christliche Werte authentisch kommuniziert."
))


# ════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════

output_path = "/Users/benjaminrumold/Downloads/cio-redesign/CiO_Redesign_InfoSheet.docx"
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Document saved to: {output_path}")
